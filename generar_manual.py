#!/usr/bin/env python3
"""
Generador del Manual de Usuario - Formatos Odontologicos v1.0
Ejecutar: python generar_manual.py
Salida: Manual_Usuario_Formatos_Odontologicos_v1.docx
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime


def configurar_estilos(doc):
    """Configura estilos base del documento."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Calibri'
        heading_style.font.color.rgb = RGBColor(0x1A, 0x56, 0x7E)
        if level == 1:
            heading_style.font.size = Pt(22)
            heading_style.paragraph_format.space_before = Pt(24)
            heading_style.paragraph_format.space_after = Pt(12)
        elif level == 2:
            heading_style.font.size = Pt(16)
            heading_style.paragraph_format.space_before = Pt(18)
            heading_style.paragraph_format.space_after = Pt(8)
        else:
            heading_style.font.size = Pt(13)
            heading_style.paragraph_format.space_before = Pt(12)
            heading_style.paragraph_format.space_after = Pt(6)


def agregar_portada(doc):
    """Crea la portada del manual."""
    for _ in range(6):
        doc.add_paragraph('')

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo.add_run('Manual de Usuario')
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x7E)
    run.bold = True

    subtitulo = doc.add_paragraph()
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitulo.add_run('Formatos Odontologicos')
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0xAB)

    doc.add_paragraph('')

    linea = doc.add_paragraph()
    linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = linea.add_run('Universidad Nacional Mayor de San Marcos')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph('')
    doc.add_paragraph('')

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run(f'Version 1.0\n{datetime.date.today().strftime("%B %Y")}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_page_break()


def agregar_captura(doc, descripcion, tamano='ancho completo'):
    """Agrega un marcador de captura de pantalla."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)

    run = p.add_run(f'[CAPTURA: {descripcion}]')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)
    run.bold = True
    run.font.italic = True

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f'Tamano sugerido: {tamano}')
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run2.font.italic = True


def agregar_nota(doc, texto):
    """Agrega una nota/tip destacado."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run('NOTA: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
    run.font.size = Pt(10)

    run2 = p.add_run(texto)
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def agregar_paso(doc, numero, texto):
    """Agrega un paso numerado."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_after = Pt(4)

    run = p.add_run(f'Paso {numero}. ')
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x7E)

    run2 = p.add_run(texto)
    run2.font.size = Pt(11)


def agregar_tabla(doc, headers, rows):
    """Agrega una tabla formateada."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph('')


def cap_presentacion(doc):
    """Capitulo 1: Presentacion."""
    doc.add_heading('1. Presentacion', level=1)

    doc.add_paragraph(
        'Formatos Odontologicos es una aplicacion de escritorio disenada para el '
        'control y gestion de clinicas odontologicas de la Universidad Nacional '
        'Mayor de San Marcos. Reemplaza los formatos manuales y hojas de calculo '
        'que se usaban anteriormente para registrar el consumo de materiales, '
        'los tratamientos realizados, la asistencia de docentes y los reportes '
        'administrativos.'
    )

    doc.add_heading('1.1 Para que sirve', level=2)
    doc.add_paragraph(
        'Este sistema permite a los asistentes de enfermeria y odontologia:'
    )
    items = [
        'Registrar tratamientos odontologicos con sus materiales consumidos',
        'Controlar la asistencia de docentes (entrada, salida, ausencias)',
        'Gestionar catalogos de materiales, docentes, pacientes y operadores',
        'Llevar el control de pagos de los tratamientos',
        'Generar reportes en Excel para el area de Administracion',
    ]
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Cm(2)

    doc.add_heading('1.2 Glosario de terminos', level=2)
    agregar_tabla(doc,
        ['Termino', 'Significado'],
        [
            ['Docente', 'Profesor o medico que atiende en la clinica'],
            ['Especialista / Operador', 'Estudiante (pregrado o posgrado) que realiza tratamientos'],
            ['Unidad de atencion', 'Estacion fisica (sillon, modulo) donde se realizan tratamientos'],
            ['Tratamiento predefinido', 'Plantilla de tratamiento con nombre, costo sugerido y materiales'],
            ['Tratamiento', 'Instancia real de atencion a un paciente, con materiales y costo'],
            ['Materiales', 'Insumos odontologicos consumidos durante tratamientos o entregados a docentes'],
            ['Grado PRE / POS', 'Pregrado (3, 4, 5) o Posgrado (R1, R2, R3) del especialista'],
        ]
    )


def cap_requisitos(doc):
    """Capitulo 2: Requisitos e Instalacion."""
    doc.add_heading('2. Requisitos e Instalacion', level=1)

    doc.add_heading('2.1 Requisitos del sistema', level=2)
    agregar_tabla(doc,
        ['Componente', 'Requisito'],
        [
            ['Sistema operativo', 'Windows 10 o Windows 11'],
            ['RAM', '4 GB minimo (8 GB recomendado)'],
            ['Disco', '500 MB de espacio libre'],
            ['Resolucion', '1280 x 720 minimo (1920 x 1080 recomendado)'],
            ['Conexion a internet', 'No requerida (funciona de forma local)'],
        ]
    )

    doc.add_heading('2.2 Instalacion', level=2)
    doc.add_paragraph('Para instalar la aplicacion, siga estos pasos:')

    agregar_paso(doc, 1, 'Ubique el archivo de instalacion (FormatosOdontologicos-Setup-x.x.x.exe).')
    agregar_paso(doc, 2, 'Haga doble clic en el archivo para iniciar el asistente de instalacion.')
    agregar_paso(doc, 3, 'Siga las instrucciones del asistente. Puede aceptar la ruta por defecto o elegir una ubicacion personalizada.')
    agregar_paso(doc, 4, 'Haga clic en "Instalar" y espere a que finalice el proceso.')
    agregar_paso(doc, 5, 'Haga clic en "Finalizar" para completar la instalacion.')

    agregar_nota(doc, 'Se creara un acceso directo en el escritorio para iniciar la aplicacion.')

    doc.add_heading('2.3 Primera ejecucion', level=2)
    doc.add_paragraph(
        'Al iniciar la aplicacion por primera vez, el sistema creara la base de datos '
        'en la carpeta de datos de la aplicacion. Este proceso puede tardar unos segundos.'
    )
    agregar_captura(doc, 'Pantalla de inicio de la aplicacion mostrando el logo y barra de carga', 'ancho completo')


def cap_primeros_pasos(doc):
    """Capitulo 3: Primeros Pasos."""
    doc.add_heading('3. Primeros Pasos', level=1)

    doc.add_heading('3.1 Seleccion de clinica', level=2)
    doc.add_paragraph(
        'Al iniciar la aplicacion, se mostrara una pantalla de seleccion de clinica. '
        'Seleccione la clinica en la que desea trabajar haciendo clic sobre ella. '
        'Esta seleccion determina todos los datos que vera en el sistema.'
    )
    agregar_captura(doc, 'Pantalla de seleccion de clinicas mostrando las clinicas disponibles', 'ancho completo')
    agregar_nota(doc, 'Puede cambiar de clinica en cualquier momento desde el selector de la barra superior.')

    doc.add_heading('3.2 Vista general de la interfaz', level=2)
    doc.add_paragraph('La interfaz principal tiene tres areas principales:')
    doc.add_paragraph('')
    items = [
        ('Barra superior', 'Muestra el nombre de la clinica seleccionada y un selector para cambiar de clinica.'),
        ('Menu lateral', 'Panel de navegacion izquierdo con acceso a todos los modulos del sistema.'),
        ('Area de contenido', 'Panel principal donde se muestra la informacion de cada modulo.'),
    ]
    for titulo, desc in items:
        p = doc.add_paragraph()
        run = p.add_run(f'{titulo}: ')
        run.bold = True
        p.add_run(desc)

    agregar_captura(doc, 'Interfaz principal con barra superior, menu lateral y area de contenido', 'ancho completo')

    doc.add_heading('3.3 Modulos disponibles', level=2)
    doc.add_paragraph('El menu lateral se divide en dos secciones:')

    doc.add_heading('Seccion Atencion', level=3)
    items_atencion = [
        ('Dashboard', 'Panel resumen con indicadores clave y graficos.'),
        ('Tratamientos', 'Gestion de tratamientos odontologicos en curso.'),
        ('Asistencia', 'Control de asistencia de docentes.'),
    ]
    for titulo, desc in items_atencion:
        p = doc.add_paragraph()
        run = p.add_run(f'{titulo}: ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('Seccion Gestion', level=3)
    items_gestion = [
        ('Catalogos', 'Gestion centralizada de materiales, docentes, pacientes, operadores y tratamientos predefinidos.'),
        ('Pagos', 'Seguimiento de pagos de tratamientos.'),
        ('Reportes', 'Generacion de reportes en Excel.'),
        ('Unidades', 'Gestion de unidades de atencion (sillones/modulos).'),
        ('Clinicas', 'Gestion de clinicas del sistema.'),
    ]
    for titulo, desc in items_gestion:
        p = doc.add_paragraph()
        run = p.add_run(f'{titulo}: ')
        run.bold = True
        p.add_run(desc)


def cap_dashboard(doc):
    """Capitulo 4: Dashboard."""
    doc.add_heading('4. Dashboard', level=1)

    doc.add_paragraph(
        'El Dashboard es el panel principal que se muestra al iniciar sesion. '
        'Proporciona una vision general rapida del estado de la clinica.'
    )

    doc.add_heading('4.1 Indicadores clave (KPIs)', level=2)
    doc.add_paragraph('En la parte superior se muestran cuatro tarjetas con indicadores:')
    agregar_tabla(doc,
        ['Indicador', 'Descripcion'],
        [
            ['Ingresos del mes', 'Total de ingresos por tratamientos en el mes actual'],
            ['Ingresos de la semana', 'Total de ingresos de la semana actual'],
            ['Tratamientos en curso', 'Cantidad de tratamientos activos (abiertos)'],
            ['Docentes hoy', 'Cantidad de docentes registrados el dia de hoy'],
        ]
    )
    agregar_captura(doc, 'Tarjetas de KPIs del Dashboard', 'ancho completo')

    doc.add_heading('4.2 Graficos', level=2)
    doc.add_paragraph('El Dashboard incluye tres graficos informativos:')

    p = doc.add_paragraph()
    run = p.add_run('Ingresos mensuales: ')
    run.bold = True
    p.add_run('Grafico de lineas que muestra la evolucion de ingresos de los ultimos 12 meses.')

    p = doc.add_paragraph()
    run = p.add_run('Tratamientos por estado: ')
    run.bold = True
    p.add_run('Grafico circular que muestra la proporcion de tratamientos abiertos, cerrados y anulados.')

    p = doc.add_paragraph()
    run = p.add_run('Top 5 materiales: ')
    run.bold = True
    p.add_run('Grafico de barras con los 5 materiales mas utilizados.')

    agregar_captura(doc, 'Graficos del Dashboard: linea, dona y barras', 'ancho completo')

    doc.add_heading('4.3 Asistencia del dia', level=2)
    doc.add_paragraph(
        'Se muestra la lista de docentes que tienen asistencia registrada el dia de hoy, '
        'con indicadores de estado (Presente, Ausente temporalmente, Finalizo).'
    )
    agregar_captura(doc, 'Seccion de asistencia del dia en el Dashboard', 'ancho completo')

    doc.add_heading('4.4 Accesos rapidos', level=2)
    doc.add_paragraph(
        'En la parte inferior se muestran tarjetas de acceso rapido a los modulos '
        'principales: Tratamientos, Asistencia, Catalogos, Reportes y Unidades.'
    )
    agregar_captura(doc, 'Tarjetas de accesos rapidos del Dashboard', 'ancho completo')


def cap_tratamientos(doc):
    """Capitulo 5: Tratamientos."""
    doc.add_heading('5. Tratamientos', level=1)

    doc.add_paragraph(
        'La seccion de Tratamientos es donde se registran y gestionan todas las '
        'atenciones odontologicas. Muestra una cuadricula de unidades de atencion '
        'con su estado actual.'
    )

    doc.add_heading('5.1 Vista general', level=2)
    doc.add_paragraph('Al acceder a Tratamientos, vera:')
    items = [
        'Un encabezado con el titulo "Tratamientos en Curso" y un boton azul "+ Nuevo tratamiento (manual)"',
        'Una cuadricula de tarjetas (StationCard), una por cada unidad de atencion registrada',
        'Cada tarjeta muestra el numero de unidad y su estado: Libre (verde) o En curso (amarillo)',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    agregar_captura(doc, 'Pagina de Tratamientos con cuadricula de unidades libres y ocupadas', 'ancho completo')

    doc.add_heading('5.2 Crear un tratamiento', level=2)

    doc.add_heading('Opcion A: Desde una unidad libre', level=3)
    doc.add_paragraph('Haga clic en una tarjeta verde (unidad libre). Se abrira el formulario de creacion con la unidad pre-seleccionada.')

    doc.add_heading('Opcion B: Manual (sin unidad)', level=3)
    doc.add_paragraph(
        'Haga clic en el boton "+ Nuevo tratamiento (manual)". Podra seleccionar '
        'una unidad de forma opcional.'
    )

    agregar_captura(doc, 'Tarjeta de unidad libre con boton "Nuevo tratamiento"', 'mitad de pagina')

    doc.add_heading('5.3 Formulario de creacion', level=2)
    doc.add_paragraph('El formulario de creacion contiene los siguientes campos:')

    agregar_tabla(doc,
        ['Campo', 'Descripcion', 'Obligatorio'],
        [
            ['Unidad', 'Unidad de atencion (pre-seleccionada si se uso una tarjeta)', 'No'],
            ['Fecha', 'Fecha del tratamiento (se completa automaticamente con hoy)', 'Si'],
            ['Paciente', 'Busque y seleccione el paciente. Puede crear uno nuevo con "+ Nuevo paciente"', 'Si'],
            ['Operador', 'Busque y seleccione el especialista. Puede crear uno nuevo con "+ Nuevo operador"', 'Si'],
            ['Tipo de tratamiento', 'Busque y seleccione un tratamiento predefinido (auto-carga costo y materiales)', 'Si'],
            ['Monto total', 'Costo del tratamiento (se auto-completa segun el tipo seleccionado)', 'Si'],
            ['Tipo', 'Comun / Continuo / Avance', 'Si'],
        ]
    )

    agregar_captura(doc, 'Formulario de creacion de tratamiento con todos los campos', 'ancho completo')

    doc.add_heading('5.3.1 Crear paciente sobre la marcha', level=3)
    doc.add_paragraph(
        'Si el paciente no existe en el catalogo, haga clic en "+ Nuevo paciente" '
        'debajo del buscador. Se abrira un formulario pequeno donde solo necesita '
        'ingresar nombres y apellidos.'
    )
    agregar_paso(doc, 1, 'Haga clic en "+ Nuevo paciente".')
    agregar_paso(doc, 2, 'Ingrese los nombres y apellidos del paciente.')
    agregar_paso(doc, 3, 'Haga clic en "Guardar y volver". El paciente se seleccionara automaticamente.')

    doc.add_heading('5.3.2 Crear operador sobre la marcha', level=3)
    doc.add_paragraph(
        'Si el operador (especialista) no existe, haga clic en "+ Nuevo operador". '
        'Ademas de nombres y apellidos, debe indicar:'
    )
    items = [
        ('Grado:', 'Pregrado (PRE) o Posgrado (POS)'),
        ('Tipo:', 'Para PRE: 3, 4 o 5. Para POS: R1, R2 o R3'),
        ('Periodo:', 'Ano academico (se completa automaticamente con el anio actual)'),
    ]
    for titulo, desc in items:
        p = doc.add_paragraph()
        run = p.add_run(titulo + ' ')
        run.bold = True
        p.add_run(desc)

    agregar_nota(doc, 'El tipo se adapta automaticamente al grado seleccionado. Si elige PRE, solo vera 3, 4, 5. Si elige POS, vera R1, R2, R3.')

    doc.add_heading('5.3.3 Tipos de tratamiento', level=3)
    agregar_tabla(doc,
        ['Tipo', 'Descripcion', 'Monto'],
        [
            ['Comun (NORMAL)', 'Tratamiento estandar con costo fijo', 'Se establece al crear'],
            ['Continuo (CONTINUO)', 'Seguimiento sin costo adicional (ej: post-operatorio)', 'Siempre S/ 0.00'],
            ['Avance (AVANCE)', 'Sesion de avance vinculada a un tratamiento padre', 'Segun tratamiento padre'],
        ]
    )

    doc.add_heading('5.3.4 Materiales predefinidos', level=3)
    doc.add_paragraph(
        'Al seleccionar un tipo de tratamiento predefinido, el sistema carga '
        'automaticamente la lista de materiales sugeridos con sus cantidades. '
        'Estos materiales aparecen en la tabla de materiales del formulario.'
    )
    doc.add_paragraph(
        'Puede modificar las cantidades, agregar materiales adicionales o '
        'eliminar materiales que no se necesiten.'
    )

    doc.add_heading('5.4 Detalle de un tratamiento', level=2)
    doc.add_paragraph(
        'Al hacer clic en una tarjeta ocupada (amarilla), se abre el panel de '
        'detalle del tratamiento con toda la informacion y las acciones disponibles.'
    )

    doc.add_heading('5.4.1 Informacion mostrada', level=3)
    items = [
        'Numero de tratamiento (ej: #5)',
        'Estado del tratamiento (ABIERTO en azul, CERRADO en verde, ANULADO en rojo)',
        'Operador, Paciente, Fecha',
        'Monto total y monto pagado',
        'Tipo de tratamiento (NORMAL, CONTINUO, AVANCE)',
        'Estado de pago (PENDIENTE, PARCIAL, PAGADO)',
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    agregar_captura(doc, 'Panel de detalle de un tratamiento abierto con toda la informacion', 'ancho completo')

    doc.add_heading('5.4.2 Gestionar materiales', level=3)
    doc.add_paragraph('Desde el detalle puede modificar los materiales del tratamiento:')
    agregar_paso(doc, 1, 'Haga clic en una tarjeta ocupada para abrir el detalle.')
    agregar_paso(doc, 2, 'En la seccion "Materiales del tratamiento", modifique las cantidades o agregue/elimine materiales.')
    agregar_paso(doc, 3, 'Haga clic en "Guardar materiales" (aparece cuando hay cambios sin guardar).')

    agregar_nota(doc, 'Si cierra el panel sin guardar, los materiales se guardan automaticamente.')

    doc.add_heading('5.4.3 Cerrar un tratamiento', level=3)
    doc.add_paragraph('Cuando el tratamiento esta finalizado:')
    agregar_paso(doc, 1, 'Abra el detalle del tratamiento.')
    agregar_paso(doc, 2, 'Haga clic en el boton verde "Cerrar tratamiento".')
    agregar_paso(doc, 3, 'El tratamiento pasara a estado CERRADO y la unidad quedara libre.')

    doc.add_heading('5.4.4 Reabrir un tratamiento', level=3)
    doc.add_paragraph(
        'Si un tratamiento esta cerrado pero necesita modificarse, puede reabrirlo:'
    )
    agregar_paso(doc, 1, 'Abra el detalle del tratamiento cerrado.')
    agregar_paso(doc, 2, 'Haga clic en el boton "Reabrir".')
    agregar_paso(doc, 3, 'El tratamiento volviera a estado ABIERTO.')

    doc.add_heading('5.4.5 Registrar pagos', level=3)
    doc.add_paragraph('Para registrar un pago de un tratamiento:')
    agregar_paso(doc, 1, 'Abra el detalle del tratamiento.')
    agregar_paso(doc, 2, 'Haga clic en "Registrar pago".')
    agregar_paso(doc, 3, 'Ingrese la fecha y el monto a abonar.')
    agregar_paso(doc, 4, 'Haga clic en "Registrar".')

    doc.add_paragraph(
        'El sistema mostrara el resumen de monto total, pagado y saldo pendiente. '
        'Los pagos parciales van actualizando el estado de pago (PENDIENTE -> PARCIAL -> PAGADO).'
    )

    agregar_captura(doc, 'Modal de registro de pago con resumen de montos', 'mitad de pagina')

    doc.add_heading('5.4.6 Anular un tratamiento', level=3)
    doc.add_paragraph('La anulacion marca un tratamiento como invalido sin eliminarlo:')
    agregar_paso(doc, 1, 'Abra el detalle del tratamiento.')
    agregar_paso(doc, 2, 'Haga clic en el boton rojo "Anular".')
    agregar_paso(doc, 3, 'Ingrese el motivo de la anulacion (obligatorio).')
    agregar_paso(doc, 4, 'Confirme con "Si, anular".')

    agregar_nota(doc, 'La anulacion es permanente. El tratamiento quedara marcado como ANULADO y no se contara en reportes futuros.')

    doc.add_heading('5.5 Estados del tratamiento', level=2)
    agregar_tabla(doc,
        ['Estado', 'Significado', 'Acciones disponibles'],
        [
            ['ABIERTO', 'Tratamiento activo en curso', 'Cerrar, Pagar, Cambiar tipo, Anular, Editar materiales'],
            ['CERRADO', 'Tratamiento finalizado', 'Reabrir, Pagar, Anular'],
            ['ANULADO', 'Tratamiento invalidado', 'Ninguna (estado final)'],
        ]
    )

    doc.add_heading('5.6 Estados de pago', level=2)
    agregar_tabla(doc,
        ['Estado', 'Significado', 'Color del badge'],
        [
            ['PENDIENTE', 'No se ha registrado ningun pago', 'Gris'],
            ['PARCIAL', 'Se ha registrado un pago parcial', 'Amarillo'],
            ['PAGADO', 'El monto total ha sido cubierto', 'Verde'],
        ]
    )


def cap_asistencia(doc):
    """Capitulo 6: Asistencia Docente."""
    doc.add_heading('6. Asistencia Docente', level=1)

    doc.add_paragraph(
        'La seccion de Asistencia permite controlar la entrada, salida y ausencias '
        'temporales de los docentes, asi como los materiales que se les entregan.'
    )

    doc.add_heading('6.1 Vista general', level=2)
    doc.add_paragraph('Al acceder a Asistencia vera tres secciones principales:')

    items = [
        ('Tabla de docentes registrados', 'Muestra todos los docentes con su estado actual del dia (Presente, Ausente temporalmente, Finalizo, Ausente).'),
        ('Busqueda manual', 'Buscador de docentes por nombre y selector de fecha.'),
        ('Registro del dia', 'Panel de detalle que aparece al seleccionar un docente.'),
    ]
    for titulo, desc in items:
        p = doc.add_paragraph()
        run = p.add_run(f'{titulo}: ')
        run.bold = True
        p.add_run(desc)

    agregar_captura(doc, 'Pagina de Asistencia con tabla de docentes y buscador', 'ancho completo')

    doc.add_heading('6.2 Registrar entrada de un docente', level=2)
    doc.add_paragraph(
        'Cuando un docente llega a la clinica, registre su asistencia:'
    )
    agregar_paso(doc, 1, 'Busque al docente en la tabla o en el buscador "Busqueda manual".')
    agregar_paso(doc, 2, 'Haga clic en "Registrar" (si el docente esta ausente) o "Ver detalle" (si ya tiene registro).')
    agregar_paso(doc, 3, 'La hora de entrada se registra automaticamente con la hora actual.')
    agregar_paso(doc, 4, 'Se cargan automaticamente los materiales predeterminados para el dia.')

    agregar_nota(doc, 'Los materiales predeterminados se configuran desde la opcion "Editar lista predeterminada". Puede restaurar la lista con "Restaurar lista predeterminada".')

    agregar_captura(doc, 'Panel de registro del dia mostrando hora de entrada y materiales predeterminados', 'ancho completo')

    doc.add_heading('6.3 Editar hora de entrada', level=2)
    doc.add_paragraph(
        'Si necesita corregir la hora de entrada, simplemente modifique el campo '
        'de hora en el panel de detalle. El cambio se guarda automaticamente.'
    )

    doc.add_heading('6.4 Gestionar materiales del dia', level=2)
    doc.add_paragraph(
        'Los materiales se muestran en una tabla con nombre y cantidad. '
        'Puede agregar, eliminar o modificar materiales durante el dia.'
    )
    agregar_paso(doc, 1, 'Seleccione el docente en la tabla de asistencia.')
    agregar_paso(doc, 2, 'En la seccion de materiales, haga clic en "+ Agregar material" para agregar uno nuevo.')
    agregar_paso(doc, 3, 'Seleccione el material del buscador y defina la cantidad.')
    agregar_paso(doc, 4, 'Para eliminar un material, haga clic en el icono de basura.')
    agregar_paso(doc, 5, 'Haga clic en "Guardar materiales" para persistir los cambios.')

    doc.add_heading('6.5 Ausencias temporales', level=2)
    doc.add_paragraph(
        'Si un docente se ausenta temporalmente (ej: para una reunion), registre la ausencia:'
    )

    doc.add_heading('Iniciar ausencia', level=3)
    agregar_paso(doc, 1, 'Abra el detalle del docente.')
    agregar_paso(doc, 2, 'Opcionalmente, ingrese un motivo en el campo "Motivo".')
    agregar_paso(doc, 3, 'Haga clic en "Iniciar ausencia". La hora de inicio se registra automaticamente.')

    doc.add_paragraph(
        'Mientras la ausencia este abierta, el boton "Registrar salida" se deshabilitara. '
        'El docente aparecera como "Ausente temporalmente" en la tabla.'
    )

    doc.add_heading('Registrar regreso', level=3)
    doc.add_paragraph(
        'Cuando el docente regrese, haga clic en "Registrar regreso" junto a la ausencia activa. '
        'La hora de fin se registra automaticamente y se calcula la duracion.'
    )

    doc.add_heading('Eliminar una ausencia', level=3)
    doc.add_paragraph(
        'Si una ausencia fue registrada por error, puede eliminarla haciendo clic en "Eliminar" '
        'junto a la ausencia cerrada.'
    )

    agregar_captura(doc, 'Panel de ausencias con tabla de periodos y botones de accion', 'ancho completo')

    doc.add_heading('6.6 Registrar salida', level=2)
    doc.add_paragraph(
        'Cuando el docente finaliza su jornada:'
    )
    agregar_paso(doc, 1, 'Asegurese de que no haya ausencias abiertas (el boton de salida estara deshabilitado si las hay).')
    agregar_paso(doc, 2, 'Haga clic en "Registrar salida". La hora se registra automaticamente.')

    doc.add_heading('Revertir una salida', level=3)
    doc.add_paragraph(
        'Si registro la salida por error, haga clic en el icono de revertir (flecha circular) '
        'junto a la hora de salida y confirme. La asistencia volvera a estar abierta.'
    )

    doc.add_heading('6.7 Editar lista de materiales predeterminados', level=2)
    doc.add_paragraph(
        'Los materiales predeterminados son los que se cargan automaticamente al '
        'registrar la entrada de un docente. Para modificar esta lista:'
    )
    agregar_paso(doc, 1, 'Abra el detalle de cualquier docente.')
    agregar_paso(doc, 2, 'Haga clic en "Editar lista predeterminada".')
    agregar_paso(doc, 3, 'Modifique la lista (agregue, elimine o cambie cantidades).')
    agregar_paso(doc, 4, 'Haga clic en "Guardar".')

    agregar_nota(doc, 'Los cambios en la lista predeterminada afectan solo a las nuevas asistencias. Las asistencias ya registradas no se modifican.')

    doc.add_heading('6.8 Cambiar de fecha', level=2)
    doc.add_paragraph(
        'Para ver o registrar asistencias de otro dia, use el selector de fecha '
        'en la seccion "Busqueda manual". El panel de detalle se cerrara y debera '
        'seleccionar un docente nuevamente.'
    )

    doc.add_heading('6.9 Anular asistencia', level=2)
    doc.add_paragraph('Si una asistencia fue registrada por error:')
    agregar_paso(doc, 1, 'Abra el detalle del docente.')
    agregar_paso(doc, 2, 'Haga clic en el boton rojo "Anular asistencia".')
    agregar_paso(doc, 3, 'Ingrese el motivo (obligatorio).')
    agregar_paso(doc, 4, 'Confirme con "Si, anular asistencia".')

    doc.add_paragraph(
        'La asistencia quedara anulada y el docente volvera a aparecer como ausente.'
    )


def cap_catalogos(doc):
    """Capitulo 7: Catalogos."""
    doc.add_heading('7. Catalogos', level=1)

    doc.add_paragraph(
        'La seccion de Catalogos permite gestionar toda la informacion de referencia '
        'del sistema: materiales, docentes, pacientes, operadores y tratamientos '
        'predefinidos.'
    )

    doc.add_heading('7.1 Organizacion por pestanas', level=2)
    doc.add_paragraph('Los catalogos se organizan en seis pestanas:')
    pestanas = [
        ('Materiales', 'Insumos odontologicos con nombre y unidad de medida'),
        ('Docentes', 'Profesores/medicos con nombre, telefono y estado'),
        ('Pacientes', 'Pacientes con nombre y apellido'),
        ('Operadores', 'Especialistas con grado (PRE/POS), tipo y periodo'),
        ('Tratamientos Predefinidos', 'Plantillas de tratamiento con costo y materiales'),
        ('Tratamientos Realizados', 'Historial de tratamientos cerrados con pagos'),
    ]
    for i, (nombre, desc) in enumerate(pestanas, 1):
        p = doc.add_paragraph()
        run = p.add_run(f'{i}. {nombre}: ')
        run.bold = True
        p.add_run(desc)

    agregar_captura(doc, 'Pagina de Catalogos con pestanas y tabla de Materiales', 'ancho completo')

    doc.add_heading('7.2 Funciones comunes', level=2)

    doc.add_heading('Buscar registros', level=3)
    doc.add_paragraph(
        'Cada pestana tiene un campo de busqueda en la parte superior. '
        'Escriba el texto a buscar y la tabla se filtrara automaticamente '
        'despues de un breve momento.'
    )

    doc.add_heading('Ordenar columnas', level=3)
    doc.add_paragraph(
        'Haga clic en el encabezado de una columna para ordenar los registros. '
        'Un clic ordena de menor a mayor, otro clic de mayor a menor, y un tercer clic '
        'elimina el ordenamiento.'
    )

    doc.add_heading('Crear un registro', level=3)
    agregar_paso(doc, 1, 'Haga clic en el boton "Nuevo [nombre del catalogo]" (verde, esquina superior derecha).')
    agregar_paso(doc, 2, 'Complete los campos del formulario.')
    agregar_paso(doc, 3, 'Haga clic en "Guardar".')

    doc.add_heading('Editar un registro', level=3)
    agregar_paso(doc, 1, 'Haga clic en el icono de lapiz (editar) en la columna "Acciones" de la tabla.')
    agregar_paso(doc, 2, 'Modifique los campos necesarios.')
    agregar_paso(doc, 3, 'Haga clic en "Guardar".')

    doc.add_heading('Eliminar un registro', level=3)
    agregar_paso(doc, 1, 'Haga clic en el icono de basura (eliminar) en la columna "Acciones".')
    agregar_paso(doc, 2, 'Confirme la eliminacion en el dialogo que aparece.')

    agregar_nota(doc, 'Algunos registros no se pueden eliminar si tienen datos asociados (ej: un material que esta siendo usado en tratamientos).')

    doc.add_heading('7.3 Gestionar Materiales', level=2)
    doc.add_paragraph('Cada material tiene:')
    items = [
        ('Nombre:', 'Identificacion del material (ej: "Guantes descartables")'),
        ('Unidad de medida:', 'Unidad en que se maneja (ej: "Caja", "Paquete", "Unidad")'),
    ]
    for titulo, desc in items:
        p = doc.add_paragraph()
        run = p.add_run(titulo + ' ')
        run.bold = True
        p.add_run(desc)

    agregar_captura(doc, 'Pestana de Materiales con boton de nuevo material y tabla', 'ancho completo')

    doc.add_heading('7.4 Gestionar Docentes', level=2)
    doc.add_paragraph('Cada docente tiene:')
    items = [
        ('Nombres:', 'Nombres completos del docente'),
        ('Apellidos:', 'Apellidos completos del docente'),
        ('Telefono:', 'Numero de contacto'),
        ('Estado:', 'Activo o Inactivo (solo al editar)'),
    ]
    for titulo, desc in items:
        p = doc.add_paragraph()
        run = p.add_run(titulo + ' ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('7.5 Gestionar Pacientes', level=2)
    doc.add_paragraph('Los pacientes solo requieren nombres y apellidos. El estado (Activo/Inactivo) se gestiona al editar.')

    doc.add_heading('7.6 Gestionar Operadores (Especialistas)', level=2)
    doc.add_paragraph('Los operadores tienen campos adicionales:')

    agregar_tabla(doc,
        ['Campo', 'Descripcion', 'Opciones'],
        [
            ['Grado', 'Nivel academico', 'Pregrado (PRE) o Posgrado (POS)'],
            ['Tipo', 'Depende del grado', 'PRE: 3, 4, 5 / POS: R1, R2, R3'],
            ['Periodo', 'Ano academico', 'Numero entero (ej: 2026)'],
            ['DNI', 'Documento de identidad (opcional)', 'Texto libre'],
        ]
    )

    agregar_nota(doc, 'El tipo se adapta automaticamente al grado. Si cambia el grado, el tipo se reinicia a la primera opcion valida.')

    agregar_captura(doc, 'Formulario de creacion de operador con cascada Grado-Tipo', 'mitad de pagina')

    doc.add_heading('7.7 Gestionar Tratamientos Predefinidos', level=2)
    doc.add_paragraph(
        'Los tratamientos predefinidos son plantillas que se usan al crear tratamientos reales. '
        'Cada uno tiene:'
    )
    items = [
        ('Nombre:', 'Nombre del tratamiento (ej: "Exodoncia simple")'),
        ('Monto sugerido:', 'Costo sugerido (se auto-completa al crear un tratamiento real)'),
        ('Materiales sugeridos:', 'Lista de materiales con cantidades que se cargan automaticamente'),
    ]
    for titulo, desc in items:
        p = doc.add_paragraph()
        run = p.add_run(titulo + ' ')
        run.bold = True
        p.add_run(desc)

    doc.add_heading('Agregar materiales a un tratamiento predefinido', level=3)
    agregar_paso(doc, 1, 'Haga clic en lapiz para editar el tratamiento predefinido.')
    agregar_paso(doc, 2, 'En la seccion "Materiales sugeridos", haga clic en "+ Agregar material".')
    agregar_paso(doc, 3, 'Seleccione el material y defina la cantidad.')
    agregar_paso(doc, 4, 'Repita para cada material necesario.')
    agregar_paso(doc, 5, 'Haga clic en "Guardar".')

    doc.add_paragraph(
        'Tambien puede expandir la fila de un tratamiento en la tabla (haga clic en la flecha) '
        'para ver los materiales asociados sin abrir el editor.'
    )

    agregar_captura(doc, 'Formulario de tratamiento predefinido con tabla de materiales sugeridos', 'ancho completo')

    doc.add_heading('7.8 Tratamientos Realizados', level=2)
    doc.add_paragraph(
        'Esta pestana muestra todos los tratamientos cerrados. A diferencia de las otras '
        'pestanas, no tiene buscador pero permite acciones avanzadas:'
    )
    items = [
        ('Editar tratamiento:', 'Modifique nombre, fecha, monto, paciente, operador o materiales de un tratamiento ya cerrado'),
        ('Registrar pago:', 'Agregue pagos parciales o totales a un tratamiento'),
        ('Anular tratamiento:', 'Marque un tratamiento como invalidado (requiere motivo)'),
    ]
    for titulo, desc in items:
        p = doc.add_paragraph()
        run = p.add_run(titulo + ' ')
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        'Tambien puede expandir cada fila para ver los materiales consumidos en ese tratamiento.'
    )

    agregar_captura(doc, 'Pestana de Tratamientos Realizados con acciones de editar, pagar y anular', 'ancho completo')


def cap_cierre(doc):
    """Capitulo 8: Cambio de Clinica y Cierre."""
    doc.add_heading('8. Cambio de Clinica y Cierre', level=1)

    doc.add_heading('8.1 Cambiar de clinica', level=2)
    doc.add_paragraph(
        'Si trabaja en varias clinicas, puede cambiar desde el selector de la barra superior:'
    )
    agregar_paso(doc, 1, 'Haga clic en el nombre de la clinica actual en la barra superior.')
    agregar_paso(doc, 2, 'Se mostrara la pantalla de seleccion de clinicas.')
    agregar_paso(doc, 3, 'Seleccione la clinica a la que desea cambiar.')

    agregar_nota(doc, 'Al cambiar de clinica, todos los datos se filtraran automaticamente. Solo vera la informacion de la clinica seleccionada.')

    doc.add_heading('8.2 Cerrar la aplicacion', level=2)
    doc.add_paragraph(
        'Para cerrar la aplicacion, haga clic en el boton de cerrar (X) de la '
        'ventana o use el menu del sistema. Los datos se guardan automaticamente.'
    )


def cap_faq(doc):
    """Capitulo 9: Preguntas Frecuentes."""
    doc.add_heading('9. Preguntas Frecuentes', level=1)

    faqs = [
        ('No encuentro un paciente/operador en el buscador',
         'Puede crear uno nuevo directamente desde el formulario de creacion de tratamientos usando los botones "+ Nuevo paciente" o "+ Nuevo operador".'),
        ('Necesito modificar un tratamiento cerrado',
         'Vaya a Catalogos > pestana "Tratamientos Realizados" y haga clic en el icono de lapiz junto al tratamiento.'),
        ('Un docente tiene materiales de mas o de menos',
         'Seleccione al docente en Asistencia, modifique los materiales y haga clic en "Guardar materiales".'),
        ('Cambié la hora de entrada de un docente por error',
         'Puede editar la hora directamente en el campo de hora de entrada. El cambio se guarda automaticamente.'),
        ('No puedo registrar la salida de un docente',
         'Verifique que no tenga ausencias abiertas. Si tiene una ausencia activa, primero registre el regreso y luego la salida.'),
        ('Necesito ver asistencias de otro dia',
         'Use el selector de fecha en la seccion "Busqueda manual" de Asistencia.'),
        ('Un material no aparece en el buscador',
         'Verifique que el material este registrado en Catalogos > Materiales y que su estado sea "Activo".'),
        ('El monto de un tratamiento Continuo se puede editar',
         'No. Los tratamientos Continuos siempre tienen monto S/ 0.00 y no es editable.'),
        ('Puedo anular un tratamiento ya cerrado',
         'Si. Tanto los tratamientos ABIERTOS como CERRADOS se pueden anular desde su detalle.'),
    ]

    for pregunta, respuesta in faqs:
        p = doc.add_paragraph()
        run = p.add_run(f'P: {pregunta}')
        run.bold = True
        p2 = doc.add_paragraph(f'R: {respuesta}')
        p2.paragraph_format.left_indent = Cm(0.5)
        p2.paragraph_format.space_after = Pt(10)


def main():
    """Funcion principal que genera el manual."""
    doc = Document()

    configurar_estilos(doc)

    configurar_estilos(doc)

    configurar_estilos(doc)

    agregar_portada(doc)

    doc.add_heading('Tabla de Contenido', level=1)
    doc.add_paragraph(
        '(Para generar la tabla de contenido automatica en Word: '
        'Referencias > Tabla de contenido > Automatica)'
    )
    doc.add_page_break()

    cap_presentacion(doc)
    doc.add_page_break()

    cap_requisitos(doc)
    doc.add_page_break()

    cap_primeros_pasos(doc)
    doc.add_page_break()

    cap_dashboard(doc)
    doc.add_page_break()

    cap_tratamientos(doc)
    doc.add_page_break()

    cap_asistencia(doc)
    doc.add_page_break()

    cap_catalogos(doc)
    doc.add_page_break()

    cap_cierre(doc)
    doc.add_page_break()

    cap_faq(doc)

    output_path = 'Manual_Usuario_Formatos_Odontologicos_v1.docx'
    doc.save(output_path)
    print(f'Manual generado exitosamente: {output_path}')
    print(f'Busque el archivo en: {output_path}')


if __name__ == '__main__':
    main()
